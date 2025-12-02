from fastapi import APIRouter, HTTPException, Response, Depends, Query
from sqlalchemy.orm import Session, joinedload
from database import get_db
from models import (
    Violation,
    License,
    Inspection,
    Business,
    ActiveStorageAttachment,
    ActiveStorageBlob,
    ViolationCodePhoto,
    DocumentTemplate,
)
from docxtpl import DocxTemplate
from docx.shared import Inches
from io import BytesIO
from typing import Optional
import os
import logging
from datetime import date

import storage
from media_service import ensure_blob_browser_safe

router = APIRouter()


def _safe_add_heading(doc: DocxTemplate, text: str, level: int = 1):
    """Add a heading, falling back to a plain paragraph when the style is missing."""
    try:
        return doc.add_heading(text, level=level)
    except KeyError:
        para = doc.add_paragraph(text)
        # Try to preserve consistent styling even if the named heading style is absent.
        try:
            para.style = doc.styles['Normal']
        except Exception:
            pass
        # Lightly emphasize higher-level headings so they stand out.
        if level <= 2:
            for run in para.runs:
                run.bold = True
        return para


def _collect_violation_photos(db: Session, violation_id: int):
    """Return violation attachments with their linked codes for templating."""
    code_links = db.query(ViolationCodePhoto).filter(ViolationCodePhoto.violation_id == violation_id).all()
    codes_by_attachment = {}
    for link in code_links:
        codes_by_attachment.setdefault(link.attachment_id, []).append(link.code_id)

    attachment_rows = (
        db.query(ActiveStorageAttachment, ActiveStorageBlob)
        .join(ActiveStorageBlob, ActiveStorageBlob.id == ActiveStorageAttachment.blob_id)
        .filter(
            ActiveStorageAttachment.record_id == violation_id,
            ActiveStorageAttachment.record_type == "Violation",
            ActiveStorageAttachment.name == "photos",
        )
        .all()
    )
    photos = []
    for attachment, blob in attachment_rows:
        if not blob:
            continue
        try:
            blob = ensure_blob_browser_safe(db, blob)
        except Exception as e:
            logging.exception(f"On-demand conversion failed for blob {getattr(blob, 'key', '?')}: {e}")
        photos.append({
            "attachment": attachment,
            "blob": blob,
            "code_ids": codes_by_attachment.get(attachment.id, []),
        })
    return photos


def _append_code_photos_to_doc(doc: DocxTemplate, violation: Violation, photos: list) -> None:
    """Append a photo evidence section grouped by code (images only)."""
    if not photos:
        return
    images_only = [
        p for p in photos if (p.get("blob") and (p["blob"].content_type or "").lower().startswith("image/"))
    ]
    if not images_only:
        return

    photos_by_code = {}
    for item in images_only:
        codes = item.get("code_ids") or []
        if not codes:
            continue
        for cid in codes:
            photos_by_code.setdefault(cid, []).append(item)

    doc.add_page_break()
    _safe_add_heading(doc, "Photo Evidence", level=1)

    for code in getattr(violation, "codes", []) or []:
        code_items = photos_by_code.get(code.id, [])
        if not code_items:
            continue
        _safe_add_heading(doc, f"{code.chapter}{'.' + code.section if code.section else ''}: {code.name}", level=2)
        for item in code_items:
            try:
                client = storage.blob_service_client.get_blob_client(
                    container=storage.CONTAINER_NAME,
                    blob=item["blob"].key,
                )
                data = client.download_blob().readall()
                doc.add_picture(BytesIO(data), width=Inches(4.5))
                if item["blob"].filename:
                    caption = doc.add_paragraph(item["blob"].filename)
                    caption.style = doc.styles['Normal']
            except Exception as e:
                logging.exception(f"Failed to embed photo for code {code.id}: {e}")


def _get_template_doc(
    db: Session,
    template_id: Optional[int],
    default_filename: str,
    expected_category: str,
) -> DocxTemplate:
    """Retrieve DocxTemplate either from DB or local file, validating category."""
    if template_id:
        template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
        if not template:
            raise HTTPException(status_code=404, detail=f"Template with id {template_id} not found")
        if template.category != expected_category:
            raise HTTPException(
                status_code=400,
                detail=f"Selected template is for '{template.category}', but expected '{expected_category}'",
            )
        return DocxTemplate(BytesIO(template.content))

    template_path = os.path.join(os.path.dirname(__file__), "../templates", default_filename)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail=f"Template not found: {default_filename}")
    return DocxTemplate(template_path)


@router.get("/violation/{violation_id}/notice")
def generate_violation_notice(
    violation_id: int,
    template_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    storage.ensure_initialized()

    violation = (
        db.query(Violation)
        .filter(Violation.id == violation_id)
        .options(
            joinedload(Violation.address),
            joinedload(Violation.codes),
            joinedload(Violation.user),
        )
        .first()
    )
    if not violation:
        raise HTTPException(status_code=404, detail="Violation not found")

    photos = _collect_violation_photos(db, violation_id)

    context = {
        "today": date.today().strftime("%m/%d/%Y"),
        "combadd": violation.address.combadd if violation.address else "",
        "premisezip": violation.address.premisezip if violation.address else "",
        "ownername": violation.address.ownername if violation.address else "",
        "owneraddress": violation.address.owneraddress if violation.address else "",
        "ownercity": violation.address.ownercity if violation.address else "",
        "ownerstate": violation.address.ownerstate if violation.address else "",
        "ownerzip": violation.address.ownerzip if violation.address else "",
        "created_at": violation.created_at.strftime("%m/%d/%Y") if violation.created_at else "",
        "deadline_date": violation.deadline_date.strftime("%m/%d/%Y") if violation.deadline_date else "",
        "userphone": violation.user.phone if violation.user and violation.user.phone else "",
        "username": violation.user.name if violation.user and violation.user.name else "",
        "violation_codes": [
            {
                "chapter": c.chapter,
                "section": c.section,
                "name": c.name,
                "description": c.description,
            }
            for c in (violation.codes or [])
        ],
    }

    doc = _get_template_doc(db, template_id, "violation_notice_template.docx", "violation")
    doc.render(context)
    _append_code_photos_to_doc(doc, violation, photos)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="violation_notice_{violation_id}.docx"'
    }
    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.get("/violation/{violation_id}/compliance-letter")
def generate_compliance_letter(
    violation_id: int,
    template_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    violation = (
        db.query(Violation)
        .filter(Violation.id == violation_id)
        .options(
            joinedload(Violation.address),
            joinedload(Violation.codes),
            joinedload(Violation.user),
        )
        .first()
    )
    if not violation:
        raise HTTPException(status_code=404, detail="Violation not found")

    address = violation.address
    today_str = date.today().strftime("%B %d, %Y")
    created_at = violation.created_at.strftime("%B %d, %Y") if violation.created_at else ""
    resolved_at = violation.updated_at.strftime("%B %d, %Y") if violation.updated_at else today_str

    owner_name = address.ownername if address and address.ownername else "Property Owner"
    owner_address_line1 = address.owneraddress if address and address.owneraddress else ""
    owner_city = address.ownercity if address and address.ownercity else ""
    owner_state = address.ownerstate if address and address.ownerstate else ""
    owner_zip = address.ownerzip if address and address.ownerzip else ""
    property_reference = address.combadd if address and address.combadd else ""

    codes_context = [
        {
            "chapter": code.chapter,
            "section": code.section,
            "name": code.name,
            "description": code.description,
        }
        for code in (violation.codes or [])
    ]

    context = {
        "today": today_str,
        "combadd": property_reference,
        "ownername": owner_name,
        "owneraddress": owner_address_line1,
        "ownercity": owner_city,
        "ownerstate": owner_state,
        "ownerzip": owner_zip,
        "created_at": created_at,
        "resolved_at": resolved_at,
        "username": (
            violation.user.name
            if violation.user and violation.user.name
            else violation.user.email if violation.user else ""
        ),
        "userphone": (violation.user.phone if violation.user and violation.user.phone else ""),
        "violation_codes": codes_context,
    }

    doc = _get_template_doc(db, template_id, "compliance_notice_template.docx", "compliance")
    doc.render(context)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="compliance_letter_{violation_id}.docx"'
    }
    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.get("/license/{license_id}/download")
def generate_license_document(
    license_id: int,
    template_id: Optional[int] = Query(None),
    db: Session = Depends(get_db),
):
    lic = (
        db.query(License)
        .filter(License.id == license_id)
        .options(joinedload(License.inspection).joinedload(Inspection.address))
        .first()
    )
    if not lic:
        raise HTTPException(status_code=404, detail="License not found")

    inspection = lic.inspection
    address = inspection.address if inspection else None

    business = None
    if getattr(lic, 'business_id', None):
        try:
            business = db.query(Business).filter(Business.id == lic.business_id).first()
        except Exception:
            business = None

    template_map = {
        1: "FY26 Business License Template.docx",
        2: "FY26 Conditional SFR License .docx",
        3: "FY26 Multi Family License Template.docx",
    }
    tpl_filename = template_map.get(lic.license_type, "FY26 Business License Template.docx")

    context = {
        "today": date.today().strftime("%m/%d/%Y"),
        "license_number": lic.license_number or "",
        "date_issued": lic.date_issued.strftime("%m/%d/%Y") if lic.date_issued else "",
        "expiration_date": lic.expiration_date.strftime("%m/%d/%Y") if lic.expiration_date else "",
        "fiscal_year": lic.fiscal_year or "",
        "conditions": lic.conditions or "",
        "combadd": address.combadd if address and getattr(address, 'combadd', None) else "",
        "premisezip": address.premisezip if address and getattr(address, 'premisezip', None) else "",
        "ownername": address.ownername if address and getattr(address, 'ownername', None) else "",
        "owneraddress": address.owneraddress if address and getattr(address, 'owneraddress', None) else "",
        "ownercity": address.ownercity if address and getattr(address, 'ownercity', None) else "",
        "ownerstate": address.ownerstate if address and getattr(address, 'ownerstate', None) else "",
        "ownerzip": address.ownerzip if address and getattr(address, 'ownerzip', None) else "",
        "business_name": business.name if business and getattr(business, 'name', None) else "",
        "business_address": getattr(business, 'address', '') if business else "",
    }

    doc = _get_template_doc(db, template_id, tpl_filename, "license")
    doc.render(context)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    def _sanitize(value: str):
        return ''.join([c if c.isalnum() else '_' for c in (value or '')]).strip('_') or f'license_{license_id}'

    primary_label = (
        business.name if business and business.name else (
            address.combadd if address and getattr(address, 'combadd', None) else f'license_{license_id}'
        )
    )
    safe = _sanitize(primary_label)

    headers = {
        'Content-Disposition': f'attachment; filename="{safe}_license_{license_id}.docx"'
    }
    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
